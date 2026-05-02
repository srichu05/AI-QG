"""
extractor.py
============
Extract raw text from uploaded documents (PDF, DOCX, TXT).
Handles encoding issues, corrupt files, and noisy extraction artefacts.
"""

import logging
from pathlib import Path

from modules.utils import clean_text, timed

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------

@timed
def extract_from_pdf(filepath: str) -> str:
    """
    Extract text from a PDF file page-by-page using PyPDF2.

    Parameters
    ----------
    filepath : str
        Absolute path to the PDF file.

    Returns
    -------
    str
        Extracted and cleaned text.
    """
    from PyPDF2 import PdfReader

    reader = PdfReader(filepath)
    pages_text: list[str] = []

    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text()
            if text:
                pages_text.append(text)
        except Exception as exc:
            logger.warning("Could not extract page %d: %s", i, exc)

    raw = "\n\n".join(pages_text)
    logger.info("PDF: extracted %d pages, %d chars", len(pages_text), len(raw))
    return clean_text(raw)


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

@timed
def extract_from_docx(filepath: str) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Parameters
    ----------
    filepath : str
        Absolute path to the DOCX file.

    Returns
    -------
    str
        Extracted and cleaned text.
    """
    from docx import Document

    doc = Document(filepath)
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    raw = "\n\n".join(paragraphs)
    logger.info("DOCX: extracted %d paragraphs, %d chars", len(paragraphs), len(raw))
    return clean_text(raw)


# ---------------------------------------------------------------------------
# TXT extraction
# ---------------------------------------------------------------------------

@timed
def extract_from_txt(filepath: str) -> str:
    """
    Read a plain text file with automatic encoding detection.

    Parameters
    ----------
    filepath : str
        Absolute path to the TXT file.

    Returns
    -------
    str
        Extracted and cleaned text.
    """
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]

    for enc in encodings:
        try:
            text = Path(filepath).read_text(encoding=enc)
            logger.info("TXT: read %d chars with encoding %s", len(text), enc)
            return clean_text(text)
        except (UnicodeDecodeError, UnicodeError):
            continue

    # Fallback: read with errors='replace'
    text = Path(filepath).read_text(encoding="utf-8", errors="replace")
    logger.warning("TXT: used fallback encoding with replacement chars")
    return clean_text(text)


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

_EXTRACTORS = {
    "pdf": extract_from_pdf,
    "docx": extract_from_docx,
    "txt": extract_from_txt,
}


def extract_text(filepath: str) -> str:
    """
    Extract text from a file based on its extension.

    Parameters
    ----------
    filepath : str
        Absolute path to the uploaded file.

    Returns
    -------
    str
        Cleaned, extracted text.

    Raises
    ------
    ValueError
        If the file extension is unsupported.
    FileNotFoundError
        If the file does not exist.
    """
    path = Path(filepath)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {filepath}")

    ext = path.suffix.lstrip(".").lower()

    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        raise ValueError(f"Unsupported file type: .{ext}")

    logger.info("Extracting text from %s (.%s)", path.name, ext)
    return extractor(filepath)
